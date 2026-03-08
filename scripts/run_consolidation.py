#!/usr/bin/env python3
"""
Çoklu extraction JSON'larından konsolide çıktılar üretir:
- Güncel Şirket Bilgileri tablosu
- Yönetim Kurulu Üyeleri tablosu
- Konsolide Esas Sözleşme

Kullanım:
  python scripts/run_consolidation.py --json-dir result/ --target-company-name "Parla Enerji" --output-dir deliverables/
  python scripts/run_consolidation.py --json-dir result/ --target-mersis 0721091699100001
"""

import argparse
import json
import sys
from pathlib import Path

# Proje kökü scripts/ bir üst
_project_root = Path(__file__).resolve().parent.parent
if str(_project_root) not in sys.path:
    sys.path.insert(0, str(_project_root))

from src.consolidation import ConsolidationService


def _filter_by_target(data: dict, name: str | None, mersis: str | None) -> dict:
    """Hedef şirket filtresi (ollama bağımlılığı olmadan script içinde)."""
    import re

    if not name and not mersis:
        return data
    companies = data.get("companies") or []
    if not isinstance(companies, list):
        return data

    def norm(t):
        if t is None or not isinstance(t, str):
            return ""
        u = t.upper().strip()
        u = re.sub(r"\s+", " ", u)
        u = u.replace("A.Ş.", "ANONIM ŞIRKETI")
        u = u.replace("\u0130", "I")  # Turkish capital İ
        return u

    nm = norm(name) if name else None
    mr = re.sub(r"\s+", "", (mersis or "").strip()) if mersis else None
    matched = []
    for c in companies:
        if not isinstance(c, dict):
            continue
        ci = c.get("company_information") or {}
        nv = (ci.get("company_name") or {}).get("value")
        mv = (ci.get("mersis_number") or {}).get("value")
        ok_name = not nm or nm in norm(nv)
        ok_mersis = not mr or (
            mv is not None and re.sub(r"\s+", "", str(mv).strip()) == mr
        )
        if ok_name and ok_mersis:
            matched.append(c)
    return {**data, "companies": matched}


def load_json(path: Path) -> dict:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="TTSG extraction sonuçlarını konsolide et"
    )
    parser.add_argument(
        "--json-dir",
        type=Path,
        required=True,
        help="Extraction JSON dosyalarının bulunduğu klasör",
    )
    parser.add_argument(
        "--target-company-name",
        type=str,
        default=None,
        help="Hedef şirket adı (örn. Parla Enerji)",
    )
    parser.add_argument(
        "--target-mersis",
        type=str,
        default=None,
        help="Hedef şirket MERSİS numarası",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=None,
        help="Çıktı klasörü (varsayılan: json-dir ile aynı)",
    )
    parser.add_argument(
        "--word",
        action="store_true",
        help="python-docx ile Word dosyası da üret",
    )
    args = parser.parse_args()

    if not args.target_company_name and not args.target_mersis:
        print(
            "En az biri gerekli: --target-company-name veya --target-mersis",
            file=sys.stderr,
        )
        sys.exit(1)

    json_dir = args.json_dir.resolve()
    if not json_dir.is_dir():
        print(f"Klasör bulunamadı: {json_dir}", file=sys.stderr)
        sys.exit(1)

    out_dir = (args.output_dir or json_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    extractions: list[dict] = []
    for path in sorted(json_dir.glob("*.json")):
        try:
            data = load_json(path)
        except Exception as e:
            print(f"Atlanıyor {path}: {e}", file=sys.stderr)
            continue

        result = data.get("result", data)
        if "companies" not in result:
            print(f"Atlanıyor {path}: 'companies' yok", file=sys.stderr)
            continue

        filtered = _filter_by_target(
            result,
            args.target_company_name,
            args.target_mersis,
        )
        companies = filtered.get("companies") or []
        if not companies:
            continue

        document_name = data.get("document_name") or path.name
        extractions.append(
            {
                "document_name": document_name,
                "result": {"companies": companies},
            }
        )

    if not extractions:
        print("Hedef şirkete ait extraction bulunamadı.", file=sys.stderr)
        sys.exit(1)

    service = ConsolidationService(extractions)
    company_info = service.company_info_table()
    board_members = service.board_members_table()
    articles = service.consolidated_articles()

    out = {
        "guncel_sirket_bilgileri": company_info,
        "yonetim_kurulu_uyeleri": board_members,
        "konsolide_esas_sozlesme": articles,
    }

    out_path = out_dir / "konsolide_sonuc.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
    print(f"Yazıldı: {out_path}")

    md_path = out_dir / "konsolide_sonuc.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("# Güncel Şirket Bilgileri\n\n")
        for k, v in company_info.items():
            f.write(f"- **{k}**: {v}\n")
        f.write("\n# Yönetim Kurulu Üyeleri\n\n")
        for row in board_members:
            f.write(
                f"- {row.get('ad_soyad_unvan', '')} | Görev bitiş: {row.get('gorev_bitis_tarihi', '')} | Kaynak: {row.get('kaynak_pdf', '')}\n"
            )
        f.write("\n# Konsolide Esas Sözleşme\n\n")
        for art in articles:
            f.write(
                f"## Madde {art.get('article_number', '')} {art.get('article_title', '') or ''}\n\n"
            )
            f.write(f"{art.get('article_text', '')}\n\n")
            f.write(f"*Kaynak: {art.get('source_ttsg', '')}*\n\n")
    print(f"Yazıldı: {md_path}")

    if args.word:
        try:
            from docx import Document

            doc = Document()
            doc.add_heading("Güncel Şirket Bilgileri", 0)
            table = doc.add_table(rows=1 + len(company_info), cols=2)
            table.style = "Table Grid"
            row = table.rows[0]
            row.cells[0].text = "Alan"
            row.cells[1].text = "Değer"
            for i, (k, v) in enumerate(company_info.items(), start=1):
                r = table.rows[i]
                r.cells[0].text = str(k)
                r.cells[1].text = str(v) if v is not None else ""
            doc.add_heading("Yönetim Kurulu Üyeleri", level=1)
            tbl2 = doc.add_table(rows=1 + len(board_members), cols=4)
            tbl2.style = "Table Grid"
            tbl2.rows[0].cells[0].text = "Ad Soyad / Unvan"
            tbl2.rows[0].cells[1].text = "Görev Bitiş"
            tbl2.rows[0].cells[2].text = "TTSG Tarih/Sayı"
            tbl2.rows[0].cells[3].text = "Kaynak PDF"
            for i, row in enumerate(board_members, start=1):
                r = tbl2.rows[i]
                r.cells[0].text = str(row.get("ad_soyad_unvan", ""))
                r.cells[1].text = str(row.get("gorev_bitis_tarihi", ""))
                r.cells[2].text = str(row.get("atandigi_ttsg_tarih_sayi", ""))
                r.cells[3].text = str(row.get("kaynak_pdf", ""))
            doc.add_heading("Konsolide Esas Sözleşme", level=1)
            for art in articles:
                doc.add_heading(
                    f"Madde {art.get('article_number', '')} {art.get('article_title', '') or ''}",
                    level=2,
                )
                doc.add_paragraph(art.get("article_text", ""))
                doc.add_paragraph(
                    f"Kaynak: {art.get('source_ttsg', '')}", style="Caption"
                )
            word_path = out_dir / "konsolide_sonuc.docx"
            doc.save(str(word_path))
            print(f"Yazıldı: {word_path}")
        except ImportError:
            print("Word çıktı için: pip install python-docx", file=sys.stderr)


if __name__ == "__main__":
    main()
